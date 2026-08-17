"""Statutory form generators — MBP-1 and DIR-8 as .docx.

Fills the exact format supplied by the user. Fields left blank in the
data are rendered as "…………………" placeholders so the professional can
complete them by hand where facts are still being gathered.
"""

import io
from datetime import datetime
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


DEFAULT_RELATIVES = [
    "Member of HUF of which Director is a member",
    "Spouse",
    "Father (including step Father)",
    "Mother (including step mother)",
    "Son (including step son)",
    "Son's wife",
    "Daughter (including step daughter)",
    "Daughter's husband",
    "Brother (including step brother)",
    "Sister (including step sister)",
]


def _blank(value: Any, dots: int = 30) -> str:
    if value in (None, "", "None"):
        return "…" * dots
    return str(value)


def _add_heading(doc: Document, text: str, size: int = 14, bold: bool = True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, underline: bool = False):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.underline = underline
    return p


def _add_line(doc: Document, text: str, bold: bool = False, size: int = 11,
              align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def _set_table_header(row, headers: List[str]):
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)


def _fy_label(fy: str) -> str:
    """'2025-26' → '01/04/2025'"""
    start = fy.split("-")[0]
    return f"01/04/{start}"


def build_mbp1(client: Dict[str, Any], director: Dict[str, Any], fy: str) -> bytes:
    """Notice of interest by director — Section 184(1) + Rule 9(1)."""
    doc = Document()

    _add_heading(doc, "FORM MBP - 1", size=14, underline=True)
    _add_heading(doc, "Notice of interest by director", size=12, bold=False)
    _add_heading(doc, "[Pursuant to section 184 (1) and rule 9(1)]", size=10, bold=False)

    doc.add_paragraph()
    _add_line(doc, "To")
    _add_line(doc, "The Board of Directors")
    _add_line(doc, client.get("name") or _blank(None))
    _add_line(doc, client.get("registered_address") or _blank(None, 60))
    doc.add_paragraph()

    _add_line(doc, "Dear Sir(s),")
    doc.add_paragraph()

    body = (
        f"I, {director.get('name') or _blank(None)}, son of "
        f"{director.get('father_name') or _blank(None)}, resident of "
        f"{director.get('address') or _blank(None, 60)}, being a director in the Company "
        "hereby give notice of my interest or concern in the following company or "
        "companies, bodies corporate, firms or other association of individuals:-"
    )
    _add_line(doc, body, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_paragraph()

    # Table I — interests
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    _set_table_header(tbl.rows[0], [
        "Sr. No.",
        "Names of the Companies /bodies corporate/ firms/ association of individuals",
        "Nature of interest or concern / Change in interest or concern",
        "Shareholding",
        "Date on which interest or concern arose / changed",
    ])
    interests = director.get("interests") or []
    if not interests:
        interests = [{}] * 3  # blank rows
    for i, item in enumerate(interests, 1):
        row = tbl.add_row().cells
        row[0].text = str(i)
        row[1].text = item.get("entity_name", "")
        row[2].text = item.get("nature", "")
        row[3].text = str(item.get("shareholding", "") or "")
        row[4].text = item.get("date", "")

    doc.add_paragraph()
    _add_line(doc, "____________________________")
    _add_line(doc, f"DIN: {director.get('din') or _blank(None, 10)}")
    _add_line(doc, f"Place: {client.get('state') or 'Bengaluru'}")
    _add_line(doc, f"Date: {_fy_label(fy)}")
    doc.add_paragraph()

    # Section II — partners
    _add_line(doc,
              "NAMES OF THE OTHER PARTNERS OF THE FIRMS IN WHICH YOU ARE A PARTNER "
              "(OR) YOUR RELATIVE IS A PARTNER, TOGETHER WITH THE NAME OF THE FIRM",
              bold=True, size=10)
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    _set_table_header(tbl2.rows[0], ["SL. NO.", "Name of the firm in which YOU ARE A PARTNER", "Name of the other Partners in the firm"])
    partners = director.get("partnerships") or []
    if not partners:
        partners = [{}]
    for i, p in enumerate(partners, 1):
        row = tbl2.add_row().cells
        row[0].text = str(i)
        row[1].text = p.get("firm_name") or "NIL"
        row[2].text = p.get("other_partners") or "NIL"

    doc.add_paragraph()

    # Section II — committees
    _add_line(doc,
              "II NAMES OF THE COMMITTEES IN WHICH YOU ARE A MEMBER OF THE OTHER COMPANY/ FIRMS",
              bold=True, size=10)
    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = "Table Grid"
    _set_table_header(tbl3.rows[0], ["SL. NO.", "COMPANY", "COMMITTEE NAME"])
    committees = director.get("committees") or []
    if not committees:
        committees = [{"company": "NIL", "name": "NIL"}]
    for i, c in enumerate(committees, 1):
        row = tbl3.add_row().cells
        row[0].text = str(i)
        row[1].text = c.get("company") or "NIL"
        row[2].text = c.get("name") or "NIL"

    doc.add_paragraph()

    # Section III — relatives
    _add_line(doc,
              "III. LIST OF RELATIVES* AS DEFINED BY SECTION 2(77) OF THE COMPANIES "
              "ACT, 2013 READ WITH COMPANIES (SPECIFICATION OF DEFINITIONS DETAILS) RULES, 2014.",
              bold=True, size=10)
    tbl4 = doc.add_table(rows=1, cols=3)
    tbl4.style = "Table Grid"
    _set_table_header(tbl4.rows[0], ["SL. NO.", "RELATIVES", "NAMES OF RELATIVES"])
    relatives_map = {r.get("relation"): r.get("name") for r in (director.get("relatives") or [])}
    for i, label in enumerate(DEFAULT_RELATIVES, 1):
        row = tbl4.add_row().cells
        row[0].text = str(i)
        row[1].text = label
        row[2].text = relatives_map.get(label) or ""

    doc.add_paragraph()
    _add_line(doc, "____________________________")
    _add_line(doc, f"DIN: {director.get('din') or _blank(None, 10)}")
    _add_line(doc, f"Place: {client.get('state') or 'Bengaluru'}")
    _add_line(doc, f"Date: {_fy_label(fy)}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_dir8(client: Dict[str, Any], director: Dict[str, Any], fy: str) -> bytes:
    """Intimation by Director (non-disqualification) — Section 164(2) + Rule 14(1)."""
    doc = Document()

    _add_heading(doc, "FORM 'DIR-8'", size=14, underline=True)
    _add_heading(doc, "Intimation by Director", size=12, bold=False)
    _add_heading(doc, "[Pursuant to Section 164(2) and rule 14(1) of Companies", size=10, bold=False)
    _add_heading(doc, "(Appointment and Qualification of Directors) Rules, 2014]", size=10, bold=False)
    doc.add_paragraph()

    reg_no = client.get("cin") or _blank(None, 20)
    nominal = client.get("nominal_capital") or _blank(None, 15)
    paid_up = client.get("paid_up_capital") or _blank(None, 15)
    _add_line(doc, f"Registration No. of Company\t: {reg_no}")
    _add_line(doc, f"Nominal Capital Rs.\t\t: {nominal}")
    _add_line(doc, f"Paid-up Capital Rs.\t\t: {paid_up}")
    _add_line(doc, f"Name of Company\t\t: {client.get('name') or ''}")
    _add_line(doc, f"Address of its Registered Office: {client.get('registered_address') or ''}")
    doc.add_paragraph()

    _add_line(doc, "To:")
    _add_line(doc, "The Board of Directors:")
    _add_line(doc, client.get("name") or "")
    doc.add_paragraph()

    body = (
        f"I, {director.get('name') or _blank(None)}, son of "
        f"{director.get('father_name') or _blank(None)}, resident of "
        f"{director.get('address') or _blank(None, 60)}, being a director in the Company "
        "hereby give notice that I am/was director in the following companies during "
        "the last three years:-"
    )
    _add_line(doc, body, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_paragraph()

    # Other directorships
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _set_table_header(tbl.rows[0], ["Name of the Company", "Date of Appointment", "Date of Cessation"])
    others = director.get("other_directorships") or []
    if not others:
        others = [{}, {}]
    for o in others:
        row = tbl.add_row().cells
        row[0].text = o.get("name", "")
        row[1].text = o.get("appointment_date", "")
        row[2].text = o.get("cessation_date", "")

    doc.add_paragraph()
    disqualified = bool(director.get("has_disqualification"))
    if disqualified:
        _add_line(doc,
                  "I further confirm that I have incurred disqualifications under section 164(2) "
                  "of the Companies Act, 2013 in the following company(s) in the previous financial "
                  "year, and that I, at present stand disqualified from being a director.",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    else:
        _add_line(doc,
                  "I further confirm that I have not incurred disqualification under section 164(2) "
                  "of the Companies Act, 2013 in any of the above companies, in the previous financial "
                  "year, and that I, at present, stand free from any disqualification from being a director.",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.add_paragraph()
    _add_line(doc, "____________________________")
    _add_line(doc, f"Name: {director.get('name') or ''}")
    _add_line(doc, f"DIN: {director.get('din') or ''}")
    _add_line(doc, f"Place: {client.get('state') or 'Bengaluru'}")
    _add_line(doc, f"Date: {_fy_label(fy)}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
